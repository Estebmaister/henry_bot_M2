"""
Document processing pipeline for RAG system.

Orchestrates document ingestion, chunking, embedding generation,
and storage with comprehensive error handling and progress tracking.
"""

import asyncio
import uuid
import logging
from docx import Document
from pathlib import Path
from typing import List, Dict, Any, Optional, Union, AsyncGenerator
from dataclasses import dataclass, field
from datetime import datetime

from .chunking import ChunkingStrategy, DocumentChunk, DocumentType, create_chunker
from .embeddings import EmbeddingService, EmbeddingResult, create_embedding_service
from .storage import VectorStore, DocumentStore, create_vector_store, create_document_store
from .chunk_store import JSONChunkStore, create_chunk_store

logger = logging.getLogger(__name__)


@dataclass
class DocumentProcessingStats:
    """Statistics for document processing operations."""
    total_documents: int = 0
    successful_documents: int = 0
    failed_documents: int = 0
    total_chunks: int = 0
    total_embeddings: int = 0
    total_processing_time_ms: float = 0.0
    total_tokens_processed: int = 0

    # Per-document type breakdown
    documents_by_type: Dict[str, int] = field(default_factory=dict)
    chunks_by_type: Dict[str, int] = field(default_factory=dict)

    # Error tracking
    errors: List[Dict[str, Any]] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "total_documents": self.total_documents,
            "successful_documents": self.successful_documents,
            "failed_documents": self.failed_documents,
            "success_rate": self.successful_documents / self.total_documents if self.total_documents > 0 else 0,
            "total_chunks": self.total_chunks,
            "total_embeddings": self.total_embeddings,
            "total_processing_time_ms": self.total_processing_time_ms,
            "total_tokens_processed": self.total_tokens_processed,
            "average_chunks_per_document": self.total_chunks / self.successful_documents if self.successful_documents > 0 else 0,
            "documents_by_type": self.documents_by_type,
            "chunks_by_type": self.chunks_by_type,
            "error_count": len(self.errors),
            "errors": self.errors[:10],  # Limit to first 10 errors
        }


@dataclass
class ProcessedDocument:
    """Result of processing a single document."""
    document_id: str
    source: str
    document_type: DocumentType
    chunks: List[DocumentChunk]
    embeddings: Optional[EmbeddingResult] = None
    processing_time_ms: float = 0.0
    success: bool = True
    error_message: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "document_id": self.document_id,
            "source": self.source,
            "document_type": self.document_type.value,
            "chunk_count": len(self.chunks),
            "embedding_generated": self.embeddings is not None,
            "processing_time_ms": self.processing_time_ms,
            "success": self.success,
            "error_message": self.error_message,
            # Limit chunks in serialization
            "chunks": [chunk.to_dict() for chunk in self.chunks[:5]],
        }


class DocumentProcessor:
    """
    High-level document processing service for RAG systems.

    Orchestrates the complete pipeline from document ingestion to
    vector storage with comprehensive error handling and monitoring.
    """

    def __init__(
        self,
        chunking_strategy: Optional[ChunkingStrategy] = None,
        embedding_service: Optional[EmbeddingService] = None,
        vector_store: Optional[VectorStore] = None,
        document_store: Optional[DocumentStore] = None,
        chunk_store: Optional[JSONChunkStore] = None,
    ):
        """Initialize document processor with services."""
        self.chunking_strategy = chunking_strategy or create_chunker()
        self.embedding_service = embedding_service or create_embedding_service()
        self.vector_store = vector_store
        self.document_store = document_store
        self.chunk_store = chunk_store

        # Processing statistics
        self.stats = DocumentProcessingStats()
        self._processing_lock = asyncio.Lock()

    async def initialize(self) -> None:
        """Initialize all required services."""
        await self.embedding_service.initialize()

        if self.vector_store:
            await self.vector_store.initialize()

        if self.document_store:
            await self.document_store.initialize()

        if self.chunk_store:
            await self.chunk_store.initialize()

        logger.info("Document processor initialized successfully")

    async def process_document(
        self,
        content: str,
        source: str,
        document_type: Optional[DocumentType] = None,
        document_id: Optional[str] = None,
        store_embeddings: bool = True,
        **chunking_kwargs
    ) -> ProcessedDocument:
        """
        Process a single document through the complete RAG pipeline.

        Args:
            content: Document text content
            source: Document source (filename, URL, etc.)
            document_type: Type of document for specialized processing
            document_id: Optional custom document ID
            store_embeddings: Whether to store embeddings in vector store
            **chunking_kwargs: Additional parameters for chunking

        Returns:
            ProcessedDocument with results and metadata
        """
        start_time = datetime.now()

        # Generate document ID if not provided
        if document_id is None:
            document_id = str(uuid.uuid4())

        # Auto-detect document type if not provided
        if document_type is None:
            document_type = self._detect_document_type(source)

        try:
            logger.info(
                f"Processing document: {source} ({document_type.value})")

            # Step 1: Chunk the document
            chunks = await self._chunk_document(
                content=content,
                document_id=document_id,
                source=source,
                document_type=document_type,
                **chunking_kwargs
            )

            if not chunks:
                raise ValueError("Document chunking produced no chunks")

            # Step 2: Generate embeddings
            embeddings = await self._generate_embeddings(chunks)

            # Step 3: Store in vector store (if available)
            if store_embeddings and self.vector_store:
                await self._store_embeddings(chunks, embeddings)

            # Step 4: Store in JSON chunk store (if available)
            if self.chunk_store:
                embedding_list = embeddings.embeddings.tolist() if embeddings else None
                await self.chunk_store.store_chunks(chunks, embedding_list)

            # Step 5: Store document metadata (if available)
            if self.document_store:
                await self._store_document_metadata(
                    document_id=document_id,
                    source=source,
                    document_type=document_type,
                    chunks=chunks
                )

            processing_time = (
                datetime.now() - start_time).total_seconds() * 1000

            result = ProcessedDocument(
                document_id=document_id,
                source=source,
                document_type=document_type,
                chunks=chunks,
                embeddings=embeddings,
                processing_time_ms=processing_time,
                success=True
            )

            # Update statistics
            await self._update_stats(result)

            logger.info(
                f"Successfully processed document: {source} - {len(chunks)} chunks")
            return result

        except Exception as e:
            processing_time = (
                datetime.now() - start_time).total_seconds() * 1000
            error_message = str(e)

            result = ProcessedDocument(
                document_id=document_id,
                source=source,
                document_type=document_type,
                chunks=[],
                processing_time_ms=processing_time,
                success=False,
                error_message=error_message
            )

            # Update error statistics
            await self._update_stats(result)

            logger.error(
                f"Failed to process document {source}: {error_message}")
            return result

    async def process_file(
        self,
        file_path: Union[str, Path],
        store_embeddings: bool = True,
        **kwargs
    ) -> ProcessedDocument:
        """
        Process a file from disk.

        Args:
            file_path: Path to the file to process
            store_embeddings: Whether to store embeddings
            **kwargs: Additional processing parameters

        Returns:
            ProcessedDocument with results
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        # Detect document type from file extension
        document_type = self._detect_document_type_from_extension(
            file_path.suffix)

        # Read file content
        content = await self._read_file_content(file_path, document_type)

        return await self.process_document(
            content=content,
            source=str(file_path),
            document_type=document_type,
            store_embeddings=store_embeddings,
            **kwargs
        )

    async def process_multiple_documents(
        self,
        documents: List[Dict[str, Any]],
        store_embeddings: bool = True,
        max_concurrent: int = 5,
        progress_callback: Optional[callable] = None
    ) -> List[ProcessedDocument]:
        """
        Process multiple documents concurrently.

        Args:
            documents: List of document dictionaries with 'content' and 'source' keys
            store_embeddings: Whether to store embeddings
            max_concurrent: Maximum number of concurrent processing tasks
            progress_callback: Optional callback for progress updates

        Returns:
            List of ProcessedDocument results
        """
        semaphore = asyncio.Semaphore(max_concurrent)

        async def process_with_semaphore(doc_data: Dict[str, Any]) -> ProcessedDocument:
            async with semaphore:
                result = await self.process_document(
                    content=doc_data["content"],
                    source=doc_data["source"],
                    document_type=doc_data.get("document_type"),
                    store_embeddings=store_embeddings,
                    **doc_data.get("chunking_kwargs", {})
                )

                if progress_callback:
                    progress_callback(result)

                return result

        # Process all documents concurrently
        tasks = [process_with_semaphore(doc) for doc in documents]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Convert exceptions to failed documents
        processed_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append(ProcessedDocument(
                    document_id=f"error_{i}",
                    source=documents[i].get("source", "unknown"),
                    document_type=DocumentType.TEXT,
                    chunks=[],
                    success=False,
                    error_message=str(result)
                ))
            else:
                processed_results.append(result)

        return processed_results

    async def _chunk_document(
        self,
        content: str,
        document_id: str,
        source: str,
        document_type: DocumentType,
        **kwargs
    ) -> List[DocumentChunk]:
        """Chunk document using the configured strategy."""
        # Run the synchronous chunking in a thread pool to avoid blocking
        def sync_chunk():
            return self.chunking_strategy.chunk_document(
                document_content=content,
                document_id=document_id,
                document_source=source,
                document_type=document_type,
                **kwargs
            )

        # Execute in thread pool to avoid blocking the event loop
        chunks = await asyncio.get_event_loop().run_in_executor(None, sync_chunk)
        return chunks

    async def _generate_embeddings(self, chunks: List[DocumentChunk]) -> EmbeddingResult:
        """Generate embeddings for document chunks."""
        chunk_texts = [chunk.text_for_embedding for chunk in chunks]
        return await self.embedding_service.embed_chunks(chunk_texts)

    async def _store_embeddings(
        self,
        chunks: List[DocumentChunk],
        embeddings: EmbeddingResult
    ) -> None:
        """Store chunks and embeddings in vector store."""
        if not self.vector_store:
            return

        # Prepare data for storage
        texts = [chunk.content for chunk in chunks]
        metadatas = [chunk.metadata.to_dict() for chunk in chunks]
        ids = [chunk.chunk_id for chunk in chunks]

        await self.vector_store.add_texts(
            texts=texts,
            embeddings=embeddings.embeddings,
            metadatas=metadatas,
            ids=ids
        )

    async def _store_document_metadata(
        self,
        document_id: str,
        source: str,
        document_type: DocumentType,
        chunks: List[DocumentChunk]
    ) -> None:
        """Store document metadata in document store."""
        if not self.document_store:
            return

        await self.document_store.store_document(
            document_id=document_id,
            source=source,
            document_type=document_type.value,
            chunk_count=len(chunks),
            metadata={
                "processed_at": datetime.now().isoformat(),
                "chunking_strategy": self.chunking_strategy.get_strategy_name(),
                "embedding_model": self.embedding_service.model.model_name,
            }
        )

    def _detect_document_type(self, source: str) -> DocumentType:
        """Auto-detect document type from source string."""
        source_lower = source.lower()

        if source_lower.endswith('.md') or source_lower.endswith('.markdown'):
            return DocumentType.MARKDOWN
        elif source_lower.endswith('.pdf'):
            return DocumentType.PDF
        elif source_lower.endswith('.docx'):
            return DocumentType.DOCX
        elif source_lower.endswith('.txt') or source_lower.endswith('.text'):
            return DocumentType.TEXT
        else:
            return DocumentType.TEXT

    def _detect_document_type_from_extension(self, extension: str) -> DocumentType:
        """Detect document type from file extension."""
        extension = extension.lower()

        if extension in ['.md', '.markdown']:
            return DocumentType.MARKDOWN
        elif extension == '.pdf':
            return DocumentType.PDF
        elif extension == '.docx':
            return DocumentType.DOCX
        elif extension in ['.txt', '.text']:
            return DocumentType.TEXT
        else:
            return DocumentType.TEXT

    async def _read_file_content(self, file_path: Path, document_type: DocumentType) -> str:
        """Read content from file based on document type."""
        if document_type == DocumentType.PDF:
            # PDF processing would require additional libraries
            raise NotImplementedError("PDF processing not yet implemented")
        elif document_type == DocumentType.DOCX:
            # Extract text from DOCX files
            try:
                doc = Document(file_path)
                full_text = []

                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text.strip())

                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            full_text.append(" | ".join(row_text))

                return "\n".join(full_text)
            except ImportError:
                raise ImportError(
                    "python-docx package is required for .docx file processing. "
                    "Install with: pip install python-docx"
                )
            except Exception as e:
                raise ValueError(
                    f"Failed to read DOCX file {file_path}: {str(e)}")
        else:
            # Simple text file reading
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # Try with different encoding
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()

    async def _update_stats(self, result: ProcessedDocument) -> None:
        """Update processing statistics."""
        async with self._processing_lock:
            self.stats.total_documents += 1

            if result.success:
                self.stats.successful_documents += 1
                self.stats.total_chunks += len(result.chunks)

                if result.embeddings:
                    self.stats.total_embeddings += result.embeddings.embeddings.shape[0]
                    self.stats.total_tokens_processed += result.embeddings.total_tokens

                # Update type-specific statistics
                doc_type = result.document_type.value
                self.stats.documents_by_type[doc_type] = self.stats.documents_by_type.get(
                    doc_type, 0) + 1
                self.stats.chunks_by_type[doc_type] = self.stats.chunks_by_type.get(
                    doc_type, 0) + len(result.chunks)

                self.stats.total_processing_time_ms += result.processing_time_ms
            else:
                self.stats.failed_documents += 1
                self.stats.errors.append({
                    "document_id": result.document_id,
                    "source": result.source,
                    "error": result.error_message,
                    "timestamp": datetime.now().isoformat(),
                })

    def get_stats(self) -> Dict[str, Any]:
        """Get comprehensive processing statistics."""
        base_stats = self.stats.to_dict()
        embedding_stats = self.embedding_service.get_stats()

        return {
            "processing_stats": base_stats,
            "embedding_stats": embedding_stats,
            "chunking_strategy": {
                "name": self.chunking_strategy.get_strategy_name(),
                "params": self.chunking_strategy.get_strategy_params(),
            }
        }

    async def reset_stats(self) -> None:
        """Reset processing statistics."""
        async with self._processing_lock:
            self.stats = DocumentProcessingStats()

    async def health_check(self) -> Dict[str, bool]:
        """Check health of all components."""
        checks = {
            "embedding_service": await self.embedding_service.health_check(),
        }

        if self.vector_store:
            checks["vector_store"] = await self.vector_store.health_check()

        if self.document_store:
            checks["document_store"] = await self.document_store.health_check()

        return checks


# Utility functions for creating document processors
def create_document_processor(**kwargs) -> DocumentProcessor:
    """
    Factory function to create document processor.

    Args:
        **kwargs: Configuration overrides

    Returns:
        Configured DocumentProcessor instance
    """
    chunking_strategy = create_chunker(**kwargs.get("chunking", {}))
    embedding_service = create_embedding_service(**kwargs.get("embedding", {}))

    # Create vector and document stores if not provided
    vector_store = kwargs.get("vector_store")
    if not vector_store:
        vector_store = create_vector_store(
            store_type=kwargs.get("vector_store_type", "faiss"),
            dimension=kwargs.get("embedding_dimension", 384),
            index_path=kwargs.get("vector_store_path",
                                  "./data/vector_store.faiss")
        )

    document_store = kwargs.get("document_store")
    if not document_store:
        document_store = create_document_store(
            store_type=kwargs.get("document_store_type", "json"),
            storage_path=kwargs.get(
                "document_store_path", "./data/documents.json")
        )

    # Create chunk store for visibility
    chunk_store = kwargs.get("chunk_store")
    if not chunk_store:
        chunk_store = create_chunk_store(
            storage_path=kwargs.get("chunk_store_path", "./data/chunks.json")
        )

    return DocumentProcessor(
        chunking_strategy=chunking_strategy,
        embedding_service=embedding_service,
        vector_store=vector_store,
        document_store=document_store,
        chunk_store=chunk_store,
    )
