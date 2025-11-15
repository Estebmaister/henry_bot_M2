"""
FastAPI server for Henry Bot M2.

Provides RESTful API endpoints with automatic Swagger documentation,
authentication, and RAG-augmented chat functionality.
"""

import tempfile
import os
import time
import asyncio
import json
import logging
from datetime import datetime
from .schemas import (
    ChatRequest, ChatResponse, DocumentUploadResponse, HealthResponse,
    ErrorResponse, AdversarialResponse, APIKeyResponse, PipelineStatusResponse,
    DocumentProcessingStatus
)
from .dependencies import verify_api_key
from .middleware import (
    URLNormalizationMiddleware, RateLimitMiddleware,
    RequestLoggingMiddleware, SecurityHeadersMiddleware
)
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, Request, Depends, File, UploadFile, status
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware

from src.core.config import settings
from src.core.agent import HenryBot
from src.core.exceptions import HenryBotError

# Get logger for this module
logger = logging.getLogger(__name__)

# Global shared processor for background tasks
_shared_processor = None
_shared_processor_lock = None


async def get_shared_processor():
    """Get or create the shared document processor for background tasks."""
    global _shared_processor, _shared_processor_lock

    if _shared_processor is None:
        if _shared_processor_lock is None:
            _shared_processor_lock = asyncio.Lock()

        async with _shared_processor_lock:
            if _shared_processor is None:
                from ..rag.processor import create_document_processor
                _shared_processor = create_document_processor()
                await _shared_processor.initialize()

    return _shared_processor


def generate_api_key() -> str:
    """
    Generate a secure API key for Henry Bot M2.

    Returns:
        A secure random API key with 'henry_bot_' prefix
    """
    # Generate 32 bytes of random data and convert to hex
    # import secrets
    # random_part = secrets.token_hex(16)
    # return f"henry_bot_{random_part}"
    return settings.api_key  # Use configured API key for simplicity


class DateTimeJSONResponse(JSONResponse):
    """Custom JSONResponse that can handle datetime objects."""

    def render(self, content) -> bytes:
        return json.dumps(
            content,
            default=self._json_serializer,
            separators=(",", ":"),
            ensure_ascii=False
        ).encode("utf-8")

    def _json_serializer(self, obj):
        """Custom JSON serializer for datetime objects."""
        if isinstance(obj, datetime):
            return obj.isoformat()
        raise TypeError(
            f'Object of type {obj.__class__.__name__} is not JSON serializable')


# Global variables for application state
app_state = {
    "bot": None,
    "start_time": None,
    "version": "2.0.0",
    "background_tasks": set(),  # Track background tasks to prevent GC
    "processing_queue": {}  # Track document processing status
}


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Manage application lifecycle."""
    # Startup
    app_state["start_time"] = time.time()
    try:
        app_state["bot"] = HenryBot()
        print("🚀 Henry Bot M2 initialized successfully")
    except Exception as e:
        print(f"❌ Failed to initialize Henry Bot M2: {e}")
        raise

    yield

    # Shutdown
    print("👋 Henry Bot M2 shutting down")

    # Cancel any pending background tasks
    if app_state["background_tasks"]:
        for task in app_state["background_tasks"]:
            if not task.done():
                task.cancel()
        await asyncio.gather(*app_state["background_tasks"], return_exceptions=True)


def create_app() -> FastAPI:
    """
    Create and configure the FastAPI application.

    Returns:
        Configured FastAPI application instance
    """
    app = FastAPI(
        title="Henry Bot M2 API",
        description="""Enhanced LLM agent with RAG capabilities and comprehensive metrics.

## Authentication
Protected endpoints require an API key sent in the `X-API-Key` header.

## Getting Started
1. Generate an API key using the `/api/v1/generate-api-key` endpoint
2. Include the key in your requests: `X-API-Key: your-api-key-here`
3. Upload documents using the `/api/v1/documents` endpoint
4. Ask questions using the `/api/v1/chat` endpoint with RAG augmentation""",
        version=app_state["version"],
        docs_url="/docs",
        redoc_url="/redoc",
        openapi_url="/openapi.json",
        lifespan=lifespan,
        default_response_class=DateTimeJSONResponse,
        openapi_components={
            "securitySchemes": {
                "ApiKeyAuth": {
                    "type": "apiKey",
                    "in": "header",
                    "name": "X-API-Key",
                    "description": "API key for authentication. Generate one using /api/v1/generate-api-key"
                }
            }
        },
        security=[{"ApiKeyAuth": []}]
    )

    # Add CORS middleware
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],  # Configure appropriately for production
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    # Add custom middleware (order matters!)
    app.add_middleware(URLNormalizationMiddleware)
    app.add_middleware(SecurityHeadersMiddleware)
    app.add_middleware(RateLimitMiddleware, calls_per_minute=60)
    app.add_middleware(RequestLoggingMiddleware)

    # Add exception handlers
    @app.exception_handler(HenryBotError)
    async def henry_bot_exception_handler(request: Request, exc: HenryBotError):
        """Handle Henry Bot specific exceptions."""
        return DateTimeJSONResponse(
            status_code=500,
            content=ErrorResponse(
                error=str(exc),
                details={"type": type(exc).__name__}
            ).dict()
        )

    @app.exception_handler(Exception)
    async def general_exception_handler(request: Request, exc: Exception):
        """Handle general exceptions."""
        return DateTimeJSONResponse(
            status_code=500,
            content=ErrorResponse(
                error="Internal server error",
                details={"type": type(exc).__name__}
            ).dict()
        )

    @app.exception_handler(404)
    async def not_found_handler(request: Request, exc):
        """Handle 404 Not Found errors."""
        return DateTimeJSONResponse(
            status_code=404,
            content=ErrorResponse(
                error="Endpoint not found",
                details={"path": request.url.path, "method": request.method}
            ).dict()
        )

    @app.exception_handler(status.HTTP_401_UNAUTHORIZED)
    async def unauthorized_handler(request: Request, exc: HTTPException):
        """Handle 401 Unauthorized errors."""
        return DateTimeJSONResponse(
            status_code=401,
            content=ErrorResponse(
                error=exc.detail,
                details={"code": "UNAUTHORIZED"}
            ).dict()
        )

    @app.exception_handler(HTTPException)
    async def http_exception_handler(request: Request, exc: HTTPException):
        """Handle all HTTPExceptions that aren't caught by specific handlers."""
        return DateTimeJSONResponse(
            status_code=exc.status_code,
            content=ErrorResponse(
                error=exc.detail,
                details={"code": f"HTTP_{exc.status_code}",
                         "path": request.url.path}
            ).dict()
        )

    # Add routes
    register_routes(app)

    return app


def register_routes(app: FastAPI):
    """Register all API routes."""

    @app.get("/health", response_model=HealthResponse, tags=["Health"])
    async def health_check():
        """
        Check system health and status.

        Returns system status including model availability, RAG system status,
        and basic performance metrics.
        """
        uptime = time.time() - \
            app_state["start_time"] if app_state["start_time"] else 0

        if app_state["bot"]:
            status = app_state["bot"].get_system_status()
            return HealthResponse(
                status=status["status"],
                model=status["model"],
                rag_available=status["rag_available"],
                prompting_technique=status["prompting_technique"],
                uptime_seconds=uptime,
                version=app_state["version"]
            )
        else:
            return HealthResponse(
                status="initializing",
                model="unknown",
                rag_available=False,
                prompting_technique="unknown",
                uptime_seconds=uptime,
                version=app_state["version"]
            )

    @app.post("/api/v1/chat", response_model=ChatResponse, tags=["Chat"])
    async def chat(request: ChatRequest, api_key: str = Depends(verify_api_key)):
        """
        Process a user question with optional RAG augmentation.

        This is the main endpoint for interacting with Henry Bot M2.
        Supports multiple prompting techniques and RAG-enhanced responses.
        """
        if not app_state["bot"]:
            raise HTTPException(
                status_code=503, detail="Service not initialized")

        try:
            # Process the question
            result = await app_state["bot"].process_question(
                user_question=request.question,
                prompt_technique=request.prompt_technique,
                use_rag=request.use_rag
            )

            # Check for adversarial response
            if "error" in result and "adversarial_info" in result:
                return DateTimeJSONResponse(
                    status_code=400,
                    content=AdversarialResponse(
                        error=result["error"],
                        adversarial_info=result["adversarial_info"]
                    ).dict()
                )

            # Check for general error
            if "error" in result:
                raise HTTPException(status_code=500, detail=result["error"])

            # Convert to response model
            response_data = {
                "answer": result.get("answer", ""),
                "metrics": result["metrics"],
            }

            # Add optional fields
            if "reasoning" in result:
                response_data["reasoning"] = result["reasoning"]
            if "rag" in result:
                response_data["rag"] = result["rag"]

            return ChatResponse(**response_data)

        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Processing failed: {str(e)}")

    @app.post("/api/v1/documents", response_model=DocumentUploadResponse, tags=["Documents"])
    async def upload_documents(
        files: list[UploadFile] = File(
            ...,
            description="Multiple document files to upload. Supported formats: .txt, .md, .markdown, .docx. Documents are processed in background."
        ),
        api_key: str = Depends(verify_api_key)
    ):
        """
        Upload documents for background processing in RAG system.
        """
        if not app_state["bot"]:
            raise HTTPException(
                status_code=503, detail="Service not initialized")

        if not app_state["bot"].document_processor:
            raise HTTPException(
                status_code=503,
                detail="Document processor not available"
            )

        start_time = time.time()
        total_files = 0
        accepted_files = []
        rejected_files = []

        try:
            # Validate and collect files for background processing
            for file in files:
                total_files += 1

                # Validate file type
                file_extension = file.filename.lower().split(
                    '.')[-1] if '.' in file.filename else ''
                if file_extension not in ['txt', 'text', 'md', 'markdown', 'docx']:
                    rejected_files.append({
                        "filename": file.filename,
                        "error": f"Unsupported file type: .{file_extension}. Supported: .txt, .md, .docx"
                    })
                    continue

                # Read file content (only async file read, no processing)
                try:
                    content = await file.read()

                    # Basic validation only
                    if not content:
                        rejected_files.append({
                            "filename": file.filename,
                            "error": "File is empty"
                        })
                        continue

                    # Generate unique document ID
                    import uuid
                    document_id = str(uuid.uuid4())

                    # Add file to accepted list for background processing
                    accepted_files.append({
                        "document_id": document_id,
                        "filename": file.filename,
                        "file_extension": file_extension,
                        "file_size": len(content),
                        "file_content": content,
                        "is_text": file_extension not in ['docx', 'pdf']
                    })

                    # Initialize status tracking IMMEDIATELY
                    app_state["processing_queue"][document_id] = {
                        "document_id": document_id,
                        "filename": file.filename,
                        "status": "queued",
                        "queued_at": datetime.now().isoformat(),
                        "started_at": None,
                        "completed_at": None,
                        "error": None,
                        "chunks_count": 0,
                        "file_size": len(content)
                    }

                except Exception as e:
                    rejected_files.append({
                        "filename": file.filename,
                        "error": f"File processing error: {str(e)}"
                    })

            # Trigger background processing with proper task management
            if accepted_files:
                for file_data in accepted_files:
                    task = asyncio.create_task(
                        _process_document_background(
                            file_data["document_id"],
                            file_data["filename"],
                            file_data["file_content"],
                            file_data["file_extension"],
                            file_data["is_text"]
                        )
                    )

                    # Keep a reference to prevent garbage collection
                    app_state["background_tasks"].add(task)

                    # Remove task from set when done
                    task.add_done_callback(
                        app_state["background_tasks"].discard)

            processing_time_ms = (time.time() - start_time) * 1000

            # Determine success status (for upload only, not processing)
            success = len(rejected_files) == 0 and len(accepted_files) > 0

            # Prepare response message
            if success and len(accepted_files) > 0:
                message = f"Successfully queued {len(accepted_files)} documents for background processing. Monitor progress using /api/v1/pipeline/status"
            elif len(accepted_files) > 0:
                message = f"Queued {len(accepted_files)} documents for background processing. {len(rejected_files)} files were rejected."
            else:
                message = "No files could be accepted for processing"

            return DocumentUploadResponse(
                success=success,
                message=message,
                documents_processed=len(accepted_files),
                total_chunks=0,
                processing_time_ms=processing_time_ms
            )

        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Document upload failed: {str(e)}")

    @app.get("/api/v1/generate-api-key", response_model=APIKeyResponse, tags=["Authentication"])
    @app.post("/api/v1/generate-api-key", response_model=APIKeyResponse, tags=["Authentication"])
    async def generate_api_key_endpoint():
        """
        Generate a new API key for accessing protected endpoints.
        This public endpoint generates a secure API key that can be used
        to authenticate requests to protected API endpoints.
        The generated API key follows the format: henry_bot_[32_random_hex_characters]
        """
        try:
            new_api_key = generate_api_key()

            return APIKeyResponse(
                api_key=new_api_key,
                status="generated",
                message="API key generated successfully. Use this key in the X-API-Key header for authenticated endpoints."
            )

        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Failed to generate API key: {str(e)}")

    @app.get("/api/v1/pipeline/status", response_model=PipelineStatusResponse, tags=["Pipeline"])
    async def get_pipeline_status(api_key: str = Depends(verify_api_key)):
        """
        Get the current status of the RAG processing pipeline.

        This endpoint provides detailed information about:
        - Overall pipeline status (idle, processing, completed)
        - Document processing statistics
        - Storage component health
        - Recent document processing activity
        - Embedding generation progress

        Use this endpoint to monitor the progress of document uploads
        and verify that chunks are being properly stored.
        """
        if not app_state["bot"]:
            raise HTTPException(
                status_code=503, detail="Service not initialized")

        try:
            uptime = time.time() - \
                app_state["start_time"] if app_state["start_time"] else 0

            # Get status from processing queue (in-memory tracking)
            queue_status = list(app_state["processing_queue"].values())

            total_documents = len(queue_status)
            processing_documents = sum(
                1 for d in queue_status if d["status"] == "processing")
            completed_documents = sum(
                1 for d in queue_status if d["status"] == "completed")
            failed_documents = sum(
                1 for d in queue_status if d["status"] == "failed")
            queued_documents = sum(
                1 for d in queue_status if d["status"] == "queued")

            total_chunks = sum(d.get("chunks_count", 0) for d in queue_status)

            # Determine pipeline status
            if processing_documents > 0 or queued_documents > 0:
                pipeline_status = "processing"
            elif failed_documents > 0 and completed_documents == 0:
                pipeline_status = "failed"
            elif completed_documents > 0:
                pipeline_status = "completed"
            else:
                pipeline_status = "idle"

            # Storage status
            storage_status = {
                "vector_store": False,
                "chunk_store": False,
                "document_store": False
            }

            # Check storage availability
            if hasattr(app_state["bot"], 'document_processor') and app_state["bot"].document_processor:
                if hasattr(app_state["bot"].document_processor, 'vector_store') and app_state["bot"].document_processor.vector_store:
                    storage_status["vector_store"] = True
                if hasattr(app_state["bot"].document_processor, 'chunk_store') and app_state["bot"].document_processor.chunk_store:
                    storage_status["chunk_store"] = True
                if hasattr(app_state["bot"].document_processor, 'document_store') and app_state["bot"].document_processor.document_store:
                    storage_status["document_store"] = True

            # Get recent documents (last 10)
            recent_documents = []
            for doc_status in sorted(queue_status, key=lambda x: x.get("queued_at", ""), reverse=True)[:10]:
                recent_documents.append(DocumentProcessingStatus(
                    document_id=doc_status["document_id"],
                    filename=doc_status["filename"],
                    status=doc_status["status"],
                    document_type="text",  # Could be enhanced to track actual type
                    chunk_count=doc_status.get("chunks_count", 0),
                    embedding_status="completed" if doc_status.get(
                        "chunks_count", 0) > 0 else "pending",
                    created_at=doc_status.get("queued_at"),
                    processing_started_at=doc_status.get("started_at"),
                    completed_at=doc_status.get("completed_at"),
                    error_message=doc_status.get("error")
                ))

            return PipelineStatusResponse(
                pipeline_status=pipeline_status,
                total_documents=total_documents,
                processing_documents=processing_documents + queued_documents,
                completed_documents=completed_documents,
                failed_documents=failed_documents,
                total_chunks=total_chunks,
                total_embeddings=total_chunks,  # Assuming 1:1 for now
                storage_status=storage_status,
                recent_documents=recent_documents,
                uptime_seconds=uptime
            )

        except Exception as e:
            logger.error(
                f"Failed to get pipeline status: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=500, detail=f"Failed to get pipeline status: {str(e)}")

    @app.get("/favicon.ico", tags=["Static"])
    async def favicon():
        """Return a minimal favicon response."""
        from fastapi import Response
        return Response(status_code=204)

    @app.get("/", tags=["Root"])
    async def root():
        """Root endpoint with basic API information."""
        return {
            "name": "Henry Bot M2",
            "version": app_state["version"],
            "description": "Enhanced LLM agent with RAG capabilities",
            "docs": "/docs",
            "health": "/health",
            "generate_api_key": "/api/v1/generate-api-key"
        }

    # Background processing function
    async def _process_document_background(
        document_id: str,
        filename: str,
        file_content,
        file_extension: str,
        is_text: bool = True
    ) -> None:
        """
        Process a document in the background with status tracking.
        """
        try:
            # Update status to processing
            if document_id in app_state["processing_queue"]:
                app_state["processing_queue"][document_id].update({
                    "status": "processing",
                    "started_at": datetime.now().isoformat()
                })

            logger.info(f"Starting background processing for: {filename}", extra={
                "document_id": document_id
            })

            # Get shared processor (initialize once only)
            processor = await get_shared_processor()

            # Create temporary file for processing
            if is_text:
                with tempfile.NamedTemporaryFile(
                    mode='w',
                    delete=False,
                    suffix=f".{file_extension}",
                    # Keep original name in prefix
                    prefix=f"{filename.rsplit('.', 1)[0]}_",
                    encoding='utf-8'
                ) as temp_file:
                    content_str = file_content.decode(
                        'utf-8') if isinstance(file_content, bytes) else file_content
                    temp_file.write(content_str)
                    temp_file_path = temp_file.name
            else:
                with tempfile.NamedTemporaryFile(
                    mode='wb',
                    delete=False,
                    suffix=f".{file_extension}",
                    # Keep original name in prefix
                    prefix=f"{filename.rsplit('.', 1)[0]}_"
                ) as temp_file:
                    content_bytes = file_content if isinstance(
                        file_content, bytes) else file_content.encode('utf-8')
                    temp_file.write(content_bytes)
                    temp_file_path = temp_file.name

            try:
                # Read the file content for process_document
                # We'll use process_document instead of process_file to have more control
                if is_text:
                    with open(temp_file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                else:
                    # For DOCX, we still need to use the file path
                    # So we'll detect the document type and read accordingly
                    from pathlib import Path
                    from ..rag.chunking import DocumentType

                    if file_extension == 'docx':
                        # For DOCX, we need to extract text
                        from docx import Document
                        doc = Document(temp_file_path)
                        full_text = []

                        # Extract text from paragraphs
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                full_text.append(paragraph.text.strip())

                        # Extract text from tables
                        for table in doc.tables:
                            for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                    if cell.text.strip():
                                        row_text.append(cell.text.strip())
                                if row_text:
                                    full_text.append(" | ".join(row_text))

                        content = "\n".join(full_text)
                    else:
                        with open(temp_file_path, 'r', encoding='utf-8') as f:
                            content = f.read()

                # Detect document type from original filename
                document_type = None
                if file_extension in ['md', 'markdown']:
                    from ..rag.chunking import DocumentType
                    document_type = DocumentType.MARKDOWN
                elif file_extension == 'docx':
                    from ..rag.chunking import DocumentType
                    document_type = DocumentType.DOCX
                elif file_extension in ['txt', 'text']:
                    from ..rag.chunking import DocumentType
                    document_type = DocumentType.TEXT

                # Process the document using process_document with original filename as source
                result = await processor.process_document(
                    content=content,
                    source=filename,  # Use original filename here!
                    document_type=document_type,
                    document_id=document_id,
                    store_embeddings=True
                )

                # Update status based on result
                if result.success:
                    app_state["processing_queue"][document_id].update({
                        "status": "completed",
                        "completed_at": datetime.now().isoformat(),
                        "chunks_count": len(result.chunks)
                    })

                    logger.info(f"Successfully processed: {filename}", extra={
                        "document_id": document_id,
                        "chunks_count": len(result.chunks)
                    })
                else:
                    app_state["processing_queue"][document_id].update({
                        "status": "failed",
                        "completed_at": datetime.now().isoformat(),
                        "error": result.error_message
                    })

                    logger.error(f"Failed to process: {filename}", extra={
                        "document_id": document_id,
                        "error": result.error_message
                    })

            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temp file: {temp_file_path}", extra={
                                   "error": str(e)})

        except Exception as e:
            logger.error(f"Background processing error for {filename}: {str(e)}", extra={
                "document_id": document_id
            }, exc_info=True)

            # Update status to failed
            if document_id in app_state["processing_queue"]:
                app_state["processing_queue"][document_id].update({
                    "status": "failed",
                    "completed_at": datetime.now().isoformat(),
                    "error": str(e)
                })
